# app.py - Python Flask 後端應用程式

from flask import Flask, request, jsonify
from flask_cors import CORS # 引入 CORS 擴展以處理跨域請求
import fitz  # PyMuPDF 用於 PDF 處理
from docx import Document # python-docx 用於 DOCX 處理
import requests # 用於發送 HTTP 請求給 Gemini API
import json # 用於處理 JSON 數據
import os # 用於讀取環境變數
import time # 用於延遲重試

# 簡化 Flask app 初始化，因為它不再服務靜態文件或模板
app = Flask(__name__)

# 將 CORS 設定為允許所有來源 (用於動態 Vercel URL)
CORS(app)

# 設定 Gemini API 金鑰
# 在生產環境中，請務必從環境變數中讀取 API 金鑰，不要硬編碼！
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "") 

# === 新增的除錯訊息：打印 GEMINI_API_KEY 的值 (使用 !r 顯示原始字串) ===
print(f"DEBUG: Backend starting. GEMINI_API_KEY (first 5 chars): {GEMINI_API_KEY[:5]}... Length: {len(GEMINI_API_KEY)}")
# =========================================================================

# 將 API_BASE_URL 設定為 Gemini API 端點
_GEMINI_BASE_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"

# === 新增的除錯訊息 (保留以確認部署版本，使用 !r 顯示原始字串) ===
print(f"DEBUG: _GEMINI_BASE_URL is set to: {_GEMINI_BASE_URL!r}")
# ===============================================

# --- 檔案解析函數 ---

def extract_text_from_pdf(file_stream):
    """從 PDF 檔案流中提取文本，嘗試保留排版。"""
    try:
        doc = fitz.open(stream=file_stream.read(), filetype="pdf")
        full_text = ""
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            # 使用 get_text("text") 提取文本，它會嘗試保留基本佈局
            text = page.get_text("text")
            full_text += text.strip() + "\n\n" # 每頁之間加兩個換行符
        doc.close()
        return full_text
    except Exception as e:
        print(f"PDF 解析錯誤: {e}")
        raise ValueError(f"無法解析 PDF 檔案: {e}")

def extract_content_from_docx(file_stream):
    """從 DOCX 檔案流中提取內容，並轉換為 HTML 格式。"""
    try:
        document = Document(file_stream)
        html_content = ""
        for paragraph in document.paragraphs:
            # 簡單地將段落轉換為 HTML <p> 標籤，保留粗體、斜體等
            html_paragraph = "<p>"
            for run in paragraph.runs:
                text = run.text
                if run.bold:
                    text = f"<strong>{text}</strong>"
                if run.italic:
                    text = f"<em>{text}</em>"
                html_paragraph += text
            html_paragraph += "</p>"
            html_content += html_paragraph + "\n"

        # 處理表格
        for table in document.tables:
            html_content += "<table>"
            for row in table.rows:
                html_content += "<tr>"
                for cell in row.cells:
                    # 簡單地將表格單元格內容轉換為 HTML <td> 標籤
                    cell_text = ""
                    for p in cell.paragraphs:
                        for run in p.runs:
                            cell_text += run.text
                    html_content += f"<td>{cell_text}</td>"
                html_content += "</tr>"
            html_content += "</table>"
            html_content += "\n" # 表格後加換行

        return html_content
    except Exception as e:
        print(f"DOCX 解析錯誤: {e}")
        raise ValueError(f"無法解析 DOCX 檔案: {e}")

# --- LLM 互動函數 ---

def call_gemini_api(cv_content):
    """呼叫 Gemini API 進行履歷驗證，包含重試邏輯。"""
    if not GEMINI_API_KEY:
        print("錯誤: GEMINI_API_KEY 環境變數未設定或為空！請確認您已在 Railway 環境中設定 GEMINI_API_KEY。")
        raise ValueError("GEMINI_API_KEY 環境變數未設定。")

    # 將提示指令和履歷內容分開，作為獨立的 text part
    instruction_text = """
        您是一個專業的履歷驗證分析師。請仔細審查以下履歷內容。
        請注意，如果這是從 DOCX 文件轉換而來，內容將是 **HTML 格式**；如果是從 PDF 轉換而來，內容將嘗試保留原始的段落和行距。
        請您根據這些格式資訊來理解內容的結構（例如：標題、列表、粗體字、表格等），並識別任何潛在的造假或盜用內容，特別是關於學歷、工作經驗、比賽或獎項、證書等。
        模擬您正在搜尋公開資訊（例如：官方網站、公開資料庫、新聞報導）以驗證這些資訊的真實性。
        如果發現任何疑點或無法核實的內容，請將其標記為「疑點」。
        對於每個疑點，請提供以下資訊：
        1. 項目名稱 (item): 履歷中被質疑的具體內容。
        2. 原因 (reason): 為什麼這是一個疑點，請簡潔說明，例如「無公開紀錄」、「資料不符」、「資訊不完整」。
        3. 組織名稱 (organizationName): 與該項目相關的官方組織或機構名稱。
        4. 組織網站 (organizationWebsite): 該組織的官方網站連結。
        5. 公開資訊連結 (publicInfoLink): 模擬找到的相關公開資訊連結（若有）。

        請**嚴格按照以下 JSON 格式**回傳結果。
        **重要提示：`hasDiscrepancies` 欄位必須是布林值，表示是否有疑點。`discrepancies` 欄位必須是一個陣列，即使沒有疑點或沒有找到具體項目，也**必須是一個空陣列**。

        範例輸出 (無疑點):
        ```json
        {
          "hasDiscrepancies": false,
          "discrepancies": []
        }
        ```

        範例輸出 (有疑點):
        ```json
        {
          "hasDiscrepancies": true,
          "discrepancies": [
            {
              "item": "國立台灣大學 電機工程學系 碩士 (2020)",
              "reason": "無公開紀錄，請提供畢業證書掃描件。",
              "organizationName": "國立台灣大學",
              "organizationWebsite": "[https://www.ntu.edu.tw/](https://www.ntu.edu.tw/)",
              "publicInfoLink": ""
            }
          ]
        }
        ```

        以下是履歷內容：
    """

    headers = {
        'Content-Type': 'application/json'
    }
    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [
                    {"text": instruction_text}, # 提示指令作為一個 text part
                    {"text": cv_content}        # 履歷內容作為另一個 text part，會被自動跳脫
                ]
            }
        ],
        "generationConfig": {
            "responseMimeType": "application/json",
            "temperature": 0.2, # 降低 temperature 以獲得更穩定的輸出
            "responseSchema": {
                "type": "OBJECT",
                "properties": {
                    "hasDiscrepancies": {"type": "BOOLEAN"},
                    "discrepancies": {
                        "type": "ARRAY",
                        "items": {
                            "type": "OBJECT",
                            "properties": {
                                "item": {"type": "STRING"},
                                "reason": {"type": "STRING"},
                                "organizationName": {"type": "STRING"},
                                "organizationWebsite": {"type": "STRING"},
                                "publicInfoLink": {"type": "STRING"}
                            }
                        }
                    }
                },
                "required": ["hasDiscrepancies", "discrepancies"]
            }
        }
    }

    max_retries = 3 # 最大重試次數
    for attempt in range(max_retries):
        try:
            # 確保 URL 和金鑰沒有前後空白字元
            base_url_clean = _GEMINI_BASE_URL.strip()
            api_key_clean = GEMINI_API_KEY.strip()
            
            full_api_url = f"{base_url_clean}?key={api_key_clean}"
            
            # === 新增的除錯打印：打印最終的 URL 字串的原始表示 ===
            print(f"DEBUG: Attempting to call Gemini API at URL: {full_api_url!r} (Attempt {attempt + 1}/{max_retries})")
            # ===============================================

            response = requests.post(
                full_api_url, # 使用清理過的 URL
                headers=headers, 
                data=json.dumps(payload),
                timeout=60 # 設置請求超時為 60 秒
            )
            response.raise_for_status() # 如果響應狀態碼不是 2xx，則引發 HTTPError
            print("Gemini API 呼叫成功。")
            return response.json() 

        except requests.exceptions.Timeout as e:
            print(f"警告: 呼叫 Gemini API 超時 (第 {attempt + 1} 次嘗試): {e}")
            if attempt < max_retries - 1:
                sleep_time = 2 ** attempt # 指數退避策略
                print(f"等待 {sleep_time} 秒後重試...")
                time.sleep(sleep_time)
            else:
                raise ConnectionError(f"無法連接到 Gemini API: 超時。 (多次嘗試失敗) {e}")
        except requests.exceptions.RequestException as e:
            error_message = f"呼叫 Gemini API 失敗 (第 {attempt + 1} 次嘗試): {e}"
            if hasattr(e, 'response') and e.response is not None:
                try:
                    error_content = e.response.json()
                    error_message += f", API回應: {error_content}"
                except json.JSONDecodeError:
                    error_content = e.response.text
                    error_message += f", API回應文本: {error_content[:200]}..."
            print(f"錯誤: {error_message}")
            if attempt < max_retries - 1:
                sleep_time = 2 ** attempt
                print(f"等待 {sleep_time} 秒後重試...")
                time.sleep(sleep_time)
            else:
                raise ConnectionError(f"無法連接到 Gemini API: {error_message} (多次嘗試失敗)")
        except json.JSONDecodeError as e:
            print(f"錯誤: 解析 Gemini API 回應失敗 (第 {attempt + 1} 次嘗試): {e}")
            raise ValueError(f"Gemini API 回應格式不正確: {e}")
    
    # 如果所有重試都失敗，理論上應該在循環中拋出異常
    raise ConnectionError("呼叫 Gemini API 失敗，所有重試均已用盡。")


# --- Flask 路由 ---

# 新增的健康檢查路由
@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({"status": "ok"}), 200

@app.route('/upload_cv', methods=['POST'])
def upload_cv():
    if 'cvFile' not in request.files:
        return jsonify({"error": "沒有檔案被上傳"}), 400

    file = request.files['cvFile']
    if file.filename == '':
        return jsonify({"error": "沒有選擇檔案"}), 400

    file_extension = file.filename.rsplit('.', 1)[1].lower()
    cv_content = ""

    try:
        if file_extension == 'txt':
            cv_content = file.stream.read().decode('utf-8')
        elif file_extension == 'pdf':
            cv_content = extract_text_from_pdf(file.stream)
        elif file_extension == 'docx':
            cv_content = extract_content_from_docx(file.stream)
        else:
            return jsonify({"error": "不支援的檔案格式，請上傳 .txt, .pdf 或 .docx 檔案"}), 400
        
        if not cv_content.strip():
            return jsonify({"error": "檔案內容為空或無法提取有效文本"}), 400

        # 將處理後的內容發送給 Gemini LLM
        llm_response = call_gemini_api(cv_content)

        # === 新增的除錯打印：打印 LLM 原始回應結構 ===
        print(f"DEBUG: Raw LLM Response: {json.dumps(llm_response, indent=2)}")
        # ===============================================

        # 這裡可以加入更多對 llm_response 結構的檢查
        # 這是修正後的程式碼，確保使用字典鍵訪問
        if isinstance(llm_response, dict) and "candidates" in llm_response and \
           len(llm_response["candidates"]) > 0 and "content" in llm_response["candidates"][0] and \
           "parts" in llm_response["candidates"][0]["content"] and \
           len(llm_response["candidates"][0]["content"]["parts"]) > 0:
            
            # 從LLM回應中提取JSON字串
            # 修正這裡的訪問方式：從 .parts 改為 ["parts"]
            response_text = llm_response["candidates"][0]["content"]["parts"][0]["text"]
            
            # === 關鍵修正：預處理 response_text，僅處理雙引號和換行符 ===
            # 移除對 \\ 的處理，因為這可能導致問題
            cleaned_response_text = response_text.replace('\\n', '\n').replace('\\"', '"')
            
            # 再次嘗試解析LLM回應的JSON字串
            # 增加一個 try-except 塊來捕獲 JSON 解析錯誤，並打印導致錯誤的字串
            try:
                parsed_llm_json = json.loads(cleaned_response_text)
            except json.JSONDecodeError as e:
                print(f"錯誤: 解析 LLM 回應的 JSON 字串失敗: {e}")
                print(f"導致錯誤的字串 (前500字元): {cleaned_response_text[:500]!r}...")
                print(f"導致錯誤的字串長度: {len(cleaned_response_text)}")
                raise ValueError(f"LLM 返回的 JSON 格式不正確: {e}")


            # 進一步驗證其內部結構是否符合responseSchema
            if "hasDiscrepancies" in parsed_llm_json and isinstance(parsed_llm_json["hasDiscrepancies"], bool) and \
               "discrepancies" in parsed_llm_json and isinstance(parsed_llm_json["discrepancies"], list):
                return jsonify(parsed_llm_json), 200
            else:
                return jsonify({"error": "LLM 回應的 JSON 結構不符預期"}), 500
        else:
            return jsonify({"error": "LLM 回應格式無效或為空"}), 500

    except ValueError as e:
        return jsonify({"error": str(e)}), 400
    except ConnectionError as e:
        return jsonify({"error": str(e)}), 503 # Service Unavailable
    except Exception as e:
        print(f"伺服器內部錯誤: {e}")
        return jsonify({"error": f"伺服器內部錯誤: {e}"}), 500

# Railway 會將請求發送到 $PORT 環境變數指定的埠號
if __name__ == '__main__':
    port = int(os.environ.get("PORT", 8080))
    app.run(host='0.0.0.0', port=port, debug=True)